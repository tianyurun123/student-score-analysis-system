"""质量分析Word导出工具函数"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO
from django.http import HttpResponse
import math
import logging

logger = logging.getLogger(__name__)


def generate_quality_analysis_doc(basic_info, analysis_texts, distribution_data, statistics):
    """生成质量分析Word文档
    
    Args:
        basic_info: 基本信息字典
        analysis_texts: 分析文本字典
        distribution_data: 成绩分布列表
        statistics: 统计信息字典
    """
    try:
        # 确保所有参数都是正确的类型
        if not isinstance(basic_info, dict):
            logger.warning(f"basic_info不是字典类型: {type(basic_info)}")
            basic_info = {}
        if not isinstance(analysis_texts, dict):
            logger.warning(f"analysis_texts不是字典类型: {type(analysis_texts)}")
            analysis_texts = {}
        if not isinstance(distribution_data, list):
            logger.warning(f"distribution_data不是列表类型: {type(distribution_data)}")
            distribution_data = []
        if not isinstance(statistics, dict):
            logger.warning(f"statistics不是字典类型: {type(statistics)}")
            statistics = {}
        
        # 调试：打印接收到的数据
        logger.info(f"generate_quality_analysis_doc接收数据 - basic_info类型: {type(basic_info)}, keys: {list(basic_info.keys()) if isinstance(basic_info, dict) else 'N/A'}")
        logger.info(f"analysis_texts类型: {type(analysis_texts)}, keys: {list(analysis_texts.keys()) if isinstance(analysis_texts, dict) else 'N/A'}")
        
        # 创建Word文档
        doc = Document()
        
        # 设置文档默认字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        doc.styles['Normal'].font.size = Pt(12)
        
        # 标题
        try:
            title = doc.add_heading('课程考试质量分析表', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if title.runs:
                title_run = title.runs[0]
                title_run.font.size = Pt(16)
                title_run.font.bold = True
            else:
                # 如果没有runs，手动添加
                title_run = title.add_run('课程考试质量分析表')
                title_run.font.size = Pt(16)
                title_run.font.bold = True
        except Exception as e:
            logger.error(f"创建标题失败: {str(e)}", exc_info=True)
            # 如果失败，使用段落替代
            title_para = doc.add_paragraph('课程考试质量分析表')
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.runs[0]
            title_run.font.size = Pt(16)
            title_run.font.bold = True
        
        # 学年学期
        try:
            semester_para = doc.add_paragraph()
            semester_run = semester_para.add_run('2025/2026学年第一学期')
            semester_run.font.size = Pt(14)
            semester_run.font.bold = True
            semester_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            logger.error(f"创建学年学期段落失败: {str(e)}", exc_info=True)
            semester_para = doc.add_paragraph('2025/2026学年第一学期')
            semester_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 基本信息表格
        info_table = doc.add_table(rows=5, cols=4)
        info_table.style = 'Light Grid Accent 1'
        
        # 设置表格列宽
        for col in info_table.columns:
            for cell in col.cells:
                cell.width = Inches(1.5)
        
        # 辅助函数：确保值是字符串（辅助内部函数）
        def _convert_to_str(value, default=''):
            """内部函数，将值转换为字符串"""
            if value is None:
                return default if default else ''
            # 处理字符串类型
            if isinstance(value, str):
                value = value.strip()
                if value == '[object Object]' or value == '[object Object]':
                    logger.warning(f"_convert_to_str检测到 '[object Object]' 字符串")
                    return default if default else ''
                return value
            # 处理字典类型
            if isinstance(value, dict):
                logger.warning(f"_convert_to_str接收到字典对象: {list(value.keys())[:5] if value else 'empty'}")
                # 如果是对象，尝试提取文本内容
                for key in ['value', 'text', 'content', 'data', 'label', 'name']:
                    if key in value:
                        result = _convert_to_str(value[key], default)
                        if result and result != '[object Object]':
                            return result
                # 如果无法提取，返回默认值而不是 "[object Object]"
                return default if default else ''
            # 处理列表或元组
            if isinstance(value, (list, tuple)):
                # 如果是列表或元组，尝试连接
                strings = [_convert_to_str(item, '') for item in value if item is not None]
                strings = [s for s in strings if s and s != '[object Object]']
                return ' '.join(strings) if strings else (default if default else '')
            # 处理数字和布尔值
            if isinstance(value, (int, float, bool)):
                return str(value)
            # 转换为字符串，确保不是 "[object Object]"
            try:
                result = str(value)
                # 如果结果是 "[object Object]"，返回默认值
                if result == '[object Object]' or result.strip() == '[object Object]' or result.startswith('<') or result.startswith('{'):
                    logger.warning(f"_convert_to_str转换结果为 '[object Object]' 或类似格式: {result[:50]}")
                    return default if default else ''
                return result
            except Exception as e:
                logger.error(f"_convert_to_str转换失败: {str(e)}, 值类型: {type(value)}")
                return default if default else ''
    
        def get_str_value(info_dict, key, default=''):
            value = info_dict.get(key, default)
            return _convert_to_str(value, default)
        
        # 填充基本信息
        info_data = [
            [('课程名称', get_str_value(basic_info, 'course_name')), ('教师姓名', get_str_value(basic_info, 'teacher_name'))],
            [('所在院(系)', get_str_value(basic_info, 'department')), ('授课班级', get_str_value(basic_info, 'class_name'))],
            [('课程学时', get_str_value(basic_info, 'hours')), ('考试人数', get_str_value(basic_info, 'exam_count', '0'))],
            [('考试性质', get_str_value(basic_info, 'exam_nature', '考试')), ('考试方法', get_str_value(basic_info, 'exam_method', '闭卷'))],
            [('考试时间', get_str_value(basic_info, 'exam_date')), ('试题来源', get_str_value(basic_info, 'question_source', '自主命题'))]
        ]
        
        for i, row_data in enumerate(info_data):
            row = info_table.rows[i]
            for j, (label, value) in enumerate(row_data):
                cell = row.cells[j * 2] if j == 0 else row.cells[j * 2 - 1]
                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                
                # 标签加粗（确保是字符串）
                if label is None:
                    label_str = ''
                elif isinstance(label, str):
                    label_str = label if label != '[object Object]' else ''
                else:
                    label_str = get_str_value({'temp': label}, 'temp', '')
                
                # 值（确保是字符串）
                if value is None:
                    value_str = ''
                elif isinstance(value, str):
                    value_str = value if value != '[object Object]' else ''
                else:
                    value_str = get_str_value({'temp': value}, 'temp', '')
                
                # 添加标签
                if label_str:
                    label_run = p.add_run(label_str)
                    label_run.font.bold = True
                    label_run.font.size = Pt(12)
                
                # 添加值
                if value_str:
                    value_run = p.add_run(value_str)
                    value_run.font.size = Pt(12)
        
        # 添加空行
        doc.add_paragraph()
        
        # 考试成绩分布表格
        dist_table = doc.add_table(rows=6, cols=3)
        dist_table.style = 'Light Grid Accent 1'
        
        # 表头
        header_row = dist_table.rows[0]
        header_cells = ['分数段', '学生数', '占总人数比例(%)']
        for i, header_text in enumerate(header_cells):
            cell = header_row.cells[i]
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header_text)
            run.font.bold = True
            run.font.size = Pt(12)
        
        # 填充分布数据
        # 确保distribution_data是列表且长度不超过表格行数
        if not isinstance(distribution_data, list):
            distribution_data = []
        
        # 表格有6行（包括表头），所以最多填充5行数据
        max_rows = min(len(distribution_data), 5)
        for i in range(max_rows):
            dist = distribution_data[i]
            if not isinstance(dist, dict):
                continue
            row = dist_table.rows[i + 1]  # i+1 因为第一行是表头
            # 确保所有值都是字符串
            label_val = get_str_value(dist, 'label', '')
            range_val = get_str_value(dist, 'range', '')
            count_val = dist.get('count', 0)
            if not isinstance(count_val, (int, float)):
                try:
                    count_val = int(count_val)
                except (ValueError, TypeError):
                    count_val = 0
            percentage_val = dist.get('percentage', 0)
            if not isinstance(percentage_val, (int, float)):
                try:
                    percentage_val = float(percentage_val)
                except (ValueError, TypeError):
                    percentage_val = 0
            
            range_text = f"{label_val}({range_val})"
            row.cells[0].text = str(range_text)
            row.cells[1].text = str(count_val)
            row.cells[2].text = f"{percentage_val:.2f}%"
            
            # 居中对齐
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.size = Pt(12)
        
        # 添加空行
        doc.add_paragraph()
        
        # 统计摘要
        stats_para = doc.add_paragraph()
        # 确保统计数据都是数字类型
        max_score = float(statistics.get('max_score', 0))
        min_score = float(statistics.get('min_score', 0))
        avg_score = float(statistics.get('avg_score', 0))
        std_dev = float(statistics.get('std_dev', 0))
        stats_text = f"最高分：{max_score:.0f}    最低分：{min_score:.0f}    平均分：{avg_score:.2f}    标准差σ：{std_dev:.2f}"
        stats_run = stats_para.add_run(stats_text)
        stats_run.font.size = Pt(12)
        
        # 添加空行
        doc.add_paragraph()
        
        # 辅助函数：确保文本是字符串（改进版，更全面地处理各种情况）
        def ensure_string(value):
            # 处理None
            if value is None:
                return ''
            
            # 处理字符串
            if isinstance(value, str):
                value = value.strip()
                # 如果是 "[object Object]"，返回空字符串
                if value == '[object Object]' or value == '[object Object]':
                    logger.warning(f"检测到 '[object Object]' 字符串，返回空字符串")
                    return ''
                return value
            
            # 处理字典（对象）
            if isinstance(value, dict):
                logger.warning(f"ensure_string接收到字典对象: {list(value.keys())[:5]}")
                # 尝试提取常见的文本字段
                for key in ['value', 'text', 'content', 'data', 'label', 'name', 'title']:
                    if key in value:
                        result = ensure_string(value[key])
                        if result:  # 如果提取到有效内容，返回
                            return result
                # 如果无法提取，尝试遍历所有值
                for k, v in value.items():
                    if isinstance(v, str) and v and v != '[object Object]':
                        return v
                # 如果都失败，返回空字符串
                return ''
            
            # 处理列表或元组
            if isinstance(value, (list, tuple)):
                # 过滤掉None和对象，只保留字符串
                strings = [ensure_string(item) for item in value if item is not None]
                # 过滤掉空字符串和"[object Object]"
                strings = [s for s in strings if s and s != '[object Object]']
                return ' '.join(strings) if strings else ''
            
            # 处理数字类型
            if isinstance(value, (int, float, bool)):
                return str(value)
            
            # 其他类型，尝试转换为字符串
            try:
                result = str(value)
                # 如果结果是 "[object Object]" 或类似格式，返回空字符串
                if result == '[object Object]' or result.strip() == '[object Object]' or result.startswith('<') or result.startswith('{'):
                    logger.warning(f"字符串转换结果为 '[object Object]' 或类似格式: {result[:50]}")
                    return ''
                return result
            except Exception as e:
                logger.error(f"ensure_string转换失败: {str(e)}, 值类型: {type(value)}")
                return ''
        
        # 试题质量分析
        qa_heading = doc.add_heading('试题质量分析', level=2)
        qa_raw = analysis_texts.get('question_quality', '')
        logger.info(f"question_quality原始值类型: {type(qa_raw)}, 值: {str(qa_raw)[:100] if qa_raw else 'None'}")
        qa_text = ensure_string(qa_raw)
        logger.info(f"question_quality转换后值: {qa_text[:100] if qa_text else 'Empty'}")
        if not qa_text or qa_text.strip() == '':
            qa_text = '（待填写）'
        qa_para = doc.add_paragraph(qa_text)
        for run in qa_para.runs:
            run.font.size = Pt(12)
        
        # 考试(卷面)成绩分析
        exam_heading = doc.add_heading('考试(卷面)成绩分析', level=2)
        exam_raw = analysis_texts.get('exam_score_analysis', '')
        logger.info(f"exam_score_analysis原始值类型: {type(exam_raw)}, 值: {str(exam_raw)[:100] if exam_raw else 'None'}")
        exam_text = ensure_string(exam_raw)
        logger.info(f"exam_score_analysis转换后值: {exam_text[:100] if exam_text else 'Empty'}")
        if not exam_text or exam_text.strip() == '':
            exam_text = '（待填写）'
        exam_para = doc.add_paragraph(exam_text)
        for run in exam_para.runs:
            run.font.size = Pt(12)
        
        # 教学效果分析及改进措施
        teaching_heading = doc.add_heading('教学效果分析及改进措施', level=2)
        teaching_raw = analysis_texts.get('teaching_effectiveness', '')
        logger.info(f"teaching_effectiveness原始值类型: {type(teaching_raw)}, 值: {str(teaching_raw)[:100] if teaching_raw else 'None'}")
        teaching_text = ensure_string(teaching_raw)
        logger.info(f"teaching_effectiveness转换后值: {teaching_text[:100] if teaching_text else 'Empty'}")
        if not teaching_text or teaching_text.strip() == '':
            teaching_text = '（待填写）'
        teaching_para = doc.add_paragraph(teaching_text)
        for run in teaching_para.runs:
            run.font.size = Pt(12)
        
        # 评价工作小组意见
        group_heading = doc.add_heading('评价工作小组意见', level=2)
        group_raw = analysis_texts.get('group_opinion', '')
        logger.info(f"group_opinion原始值类型: {type(group_raw)}, 值: {str(group_raw)[:100] if group_raw else 'None'}")
        group_text = ensure_string(group_raw)
        logger.info(f"group_opinion转换后值: {group_text[:100] if group_text else 'Empty'}")
        if not group_text or group_text.strip() == '':
            group_text = '（待填写）'
        group_para = doc.add_paragraph(group_text)
        for run in group_para.runs:
            run.font.size = Pt(12)
        
        # 负责人签字和日期
        doc.add_paragraph()
        sign_para = doc.add_paragraph()
        signer = ensure_string(analysis_texts.get('signer', ''))
        sign_date = ensure_string(analysis_texts.get('sign_date', ''))
        sign_text = f"负责人签字：{signer}    日期：{sign_date}"
        sign_run = sign_para.add_run(sign_text)
        sign_run.font.size = Pt(12)
        
    except Exception as e:
        logger.error(f"生成Word文档过程中出现异常: {str(e)}", exc_info=True)
        # 如果出现异常，创建一个包含错误信息的简单文档
        doc = Document()
        error_para = doc.add_paragraph(f'生成文档时出现错误: {str(e)}')
        return doc
    
    return doc
